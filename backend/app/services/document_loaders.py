"""Extract plain text from supported upload formats."""

from __future__ import annotations

import csv
import io
import os
from pathlib import Path

from pypdf import PdfReader
from pypdf.errors import PdfStreamError

from app.core.supported_uploads import ALLOWED_EXTENSIONS


class DocumentLoadError(Exception):
    """Raised when text cannot be extracted from an uploaded file."""


def _normalize_text(text: str) -> str:
    cleaned = (text or "").replace("\x00", " ").strip()
    if not cleaned:
        raise DocumentLoadError("No readable text was found in the file.")
    return cleaned


def load_pdf(file_path: str) -> str:
    try:
        reader = PdfReader(file_path)
        parts: list[str] = []
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                parts.append(extracted)
        return _normalize_text("\n".join(parts))
    except PdfStreamError as exc:
        raise DocumentLoadError("Invalid or corrupted PDF file.") from exc
    except Exception as exc:
        raise DocumentLoadError(f"PDF processing failed: {exc}") from exc


def load_txt(file_path: str) -> str:
    raw = Path(file_path).read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return _normalize_text(raw.decode(encoding))
        except UnicodeDecodeError:
            continue
    raise DocumentLoadError("Could not decode text file. Use UTF-8 encoding.")


def load_docx(file_path: str) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentLoadError(
            "DOCX support is unavailable. Install python-docx on the server."
        ) from exc

    try:
        document = Document(file_path)
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return _normalize_text("\n".join(parts))
    except Exception as exc:
        raise DocumentLoadError(f"DOCX processing failed: {exc}") from exc


def load_csv(file_path: str) -> str:
    try:
        raw = Path(file_path).read_text(encoding="utf-8-sig")
        reader = csv.reader(io.StringIO(raw))
        rows = [" | ".join(cell.strip() for cell in row if cell.strip()) for row in reader]
        rows = [row for row in rows if row]
        return _normalize_text("\n".join(rows))
    except Exception as exc:
        raise DocumentLoadError(f"CSV processing failed: {exc}") from exc


def load_xlsx(file_path: str) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise DocumentLoadError(
            "Excel support is unavailable. Install openpyxl on the server."
        ) from exc

    try:
        workbook = load_workbook(file_path, read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in workbook.worksheets:
            parts.append(f"# Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if cells:
                    parts.append(" | ".join(cells))
        workbook.close()
        return _normalize_text("\n".join(parts))
    except Exception as exc:
        raise DocumentLoadError(f"Excel processing failed: {exc}") from exc


def load_xls(file_path: str) -> str:
    try:
        import xlrd
    except ImportError as exc:
        raise DocumentLoadError(
            "Legacy Excel (.xls) support is unavailable. Install xlrd on the server."
        ) from exc

    try:
        workbook = xlrd.open_workbook(file_path)
        parts: list[str] = []
        for sheet in workbook.sheets():
            parts.append(f"# Sheet: {sheet.name}")
            for row_idx in range(sheet.nrows):
                cells = [
                    str(sheet.cell_value(row_idx, col_idx)).strip()
                    for col_idx in range(sheet.ncols)
                    if str(sheet.cell_value(row_idx, col_idx)).strip()
                ]
                if cells:
                    parts.append(" | ".join(cells))
        return _normalize_text("\n".join(parts))
    except Exception as exc:
        raise DocumentLoadError(f"Excel (.xls) processing failed: {exc}") from exc


def load_document(file_path: str) -> str:
    extension = os.path.splitext(file_path)[1].lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise DocumentLoadError(
            f"Unsupported file type '{extension}'. "
            f"Supported types: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    if extension == ".pdf":
        return load_pdf(file_path)
    if extension == ".txt":
        return load_txt(file_path)
    if extension == ".docx":
        return load_docx(file_path)
    if extension == ".csv":
        return load_csv(file_path)
    if extension == ".xlsx":
        return load_xlsx(file_path)
    if extension == ".xls":
        return load_xls(file_path)

    raise DocumentLoadError(f"Unsupported file type '{extension}'.")
